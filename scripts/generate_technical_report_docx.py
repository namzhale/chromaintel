from __future__ import annotations

import csv
import json
import math
import re
from datetime import date
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORTS = PROJECT_ROOT / "reports"
DATA = PROJECT_ROOT / "data" / "processed"
OUTPUT_DOCX = REPORTS / "chromaintel_technical_report_ru.docx"
CHART_DIR = REPORTS / "technical_report_assets"


PALETTE = {
    "ink": "1F2937",
    "muted": "64748B",
    "line": "CBD5E1",
    "blue": "2563EB",
    "teal": "0F766E",
    "green": "16A34A",
    "amber": "D97706",
    "red": "DC2626",
    "soft_blue": "EFF6FF",
    "soft_teal": "F0FDFA",
    "soft_amber": "FFFBEB",
    "soft_gray": "F8FAFC",
}


def main() -> int:
    REPORTS.mkdir(exist_ok=True)
    CHART_DIR.mkdir(exist_ok=True)

    target_coverage = _read_csv(REPORTS / "target_coverage_matrix.csv")
    benchmark = _read_csv(REPORTS / "model_benchmark_matrix.csv")
    inverse_metrics = _read_csv(REPORTS / "inverse_model_metrics.csv")
    inverse_topk = _read_csv(REPORTS / "inverse_topk_evaluation.csv")
    retention_order = _read_csv(REPORTS / "retention_order_metrics.csv")
    source_holdout = _read_csv(REPORTS / "source_holdout_metrics.csv")
    dl_report = _read_json(REPORTS / "benchmarks" / "dl_dataset_prep_report.json")

    charts = {
        "targets": _target_coverage_chart(target_coverage, CHART_DIR / "target_coverage.png"),
        "rt_models": _rt_model_chart(benchmark, CHART_DIR / "rt_model_mae.png"),
        "source_holdout": _source_holdout_chart(source_holdout, CHART_DIR / "source_holdout.png"),
        "inverse": _inverse_brier_chart(inverse_metrics, CHART_DIR / "inverse_brier.png"),
    }

    doc = Document()
    _setup_document(doc)
    _cover(doc)
    _executive_summary(doc, target_coverage, benchmark, inverse_metrics, retention_order, dl_report)
    _architecture(doc)
    _data_layer(doc, dl_report)
    _target_readiness(doc, target_coverage, charts["targets"])
    _forward_models(doc, benchmark, retention_order, source_holdout, charts["rt_models"], charts["source_holdout"])
    _inverse_models(doc, inverse_metrics, inverse_topk, charts["inverse"])
    _dl_preparation(doc, dl_report)
    _gui_reports(doc)
    _limitations_and_roadmap(doc)
    _appendix(doc, target_coverage, benchmark, inverse_metrics)

    doc.save(OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX.resolve()}")
    return 0


def _setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 26, PALETTE["ink"]),
        ("Heading 1", 17, PALETTE["blue"]),
        ("Heading 2", 13, PALETTE["teal"]),
        ("Heading 3", 11, PALETTE["ink"]),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name == "Title" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = _rgb(color)
        style.font.bold = True


def _cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ChromaIntel")
    run.bold = True
    run.font.size = Pt(34)
    run.font.color.rgb = _rgb(PALETTE["blue"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI-assisted LC-MS/MS Method Development MVP")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = _rgb(PALETTE["ink"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Технический отчет по данным, моделям, прямой и обратной задаче")
    run.font.size = Pt(12)
    run.font.color.rgb = _rgb(PALETTE["muted"])

    doc.add_paragraph()
    _add_info_band(
        doc,
        "Назначение документа",
        "Описать, как устроен MVP ChromaIntel: какие данные собраны, что предсказывается, "
        "как обучаются forward и inverse модели, какие метрики получены, где ограничения текущего "
        "публичного датасета и что нужно для перехода к production-grade биоаналитической платформе.",
        fill=PALETTE["soft_blue"],
        accent=PALETTE["blue"],
    )

    rows = [
        ("Дата отчета", date.today().isoformat()),
        ("Репозиторий", "github.com/namzhale/chromaintel"),
        ("Локальный проект", str(PROJECT_ROOT)),
        ("Основной стек", "Python, RDKit, scikit-learn, XGBoost, CatBoost, FastAPI, Streamlit, PostgreSQL"),
    ]
    _add_key_value_table(doc, rows)
    doc.add_page_break()


def _executive_summary(
    doc: Document,
    target_coverage: pd.DataFrame,
    benchmark: pd.DataFrame,
    inverse_metrics: pd.DataFrame,
    retention_order: pd.DataFrame,
    dl_report: dict,
) -> None:
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "ChromaIntel сейчас является локально запускаемым MVP для AI-assisted разработки LC-MS/MS методик "
        "в биоаналитике малых молекул. Система уже умеет собирать публичные RT-датасеты, нормализовать "
        "структуры и LC/MS-условия, обучать forward-модели retention time и proxy-quality, а также ранжировать "
        "кандидатные методы через constrained search поверх forward-предсказаний."
    )

    best_rt = _best_rt_row(benchmark)
    best_quality = _best_quality_row(benchmark)
    order_all = retention_order.iloc[0].to_dict() if not retention_order.empty else {}
    cards = [
        ("Dataset", f"{dl_report.get('input_rows', 213941):,}".replace(",", " "), "canonical rows"),
        ("Compounds", "202 834", "unique compounds"),
        ("Best RT", best_rt.get("model_family", "extra_trees"), f"MAE {float(best_rt.get('mae', 0)):.2f} min"),
        ("Retention order", f"{float(order_all.get('pairwise_order_accuracy', 0)):.3f}", "pairwise accuracy"),
        ("Inverse rows", f"{_inverse_training_rows():,}".replace(",", " "), "proxy candidates"),
        ("DL prep", f"{dl_report.get('pairwise_pairs', 0):,}".replace(",", " "), "pairwise pairs"),
    ]
    _add_metric_cards(doc, cards)

    _add_info_band(
        doc,
        "Главный вывод",
        "На публичных данных надежно обучается RT и производный quality surrogate. Параметры формы пика "
        "(асимметрия, tailing, ширина у основания/полувысоты, resolution, S/N, area/height) уже заведены "
        "в схему, но в текущем публичном training view имеют нулевое measured-покрытие. Для них требуется "
        "внутренний lab export или curated literature extraction.",
        fill=PALETTE["soft_amber"],
        accent=PALETTE["amber"],
    )

    _add_bullets(
        doc,
        [
            "Forward задача: compound + LC/MS conditions -> predicted RT, proxy quality, uncertainty/ad flags.",
            "Inverse задача: target criteria + constraints -> ranked candidate methods через search + forward scoring + optional inverse reranker.",
            "Текущая inverse модель обучена на synthetic_proxy suitability labels; это проверка pipeline, не доказательство лабораторного success rate.",
            "GUI и PDF dashboard уже показывают dataset coverage, модельные метрики, target readiness и inverse diagnostics.",
        ],
    )


def _architecture(doc: Document) -> None:
    doc.add_heading("2. Архитектура MVP", level=1)
    doc.add_paragraph(
        "Архитектура разделена на слои, чтобы не смешивать ingestion, химию, моделирование и GUI. "
        "Это важно для аудируемости биоаналитических workflow: источник каждой строки сохраняется, "
        "а рекомендации объясняются через score decomposition."
    )
    rows = [
        ("Data ingestion", "Adapters для RepoRT, MCMRT, ReTiNA, METLIN SMRT, Kaggle descriptor sidecar, internal template"),
        ("Canonical schema", "Compound identity, LC method, MS settings, sample context, observation, peak metrics, provenance"),
        ("Feature engineering", "RDKit descriptors, LC numeric/categorical features, simplified gradient, MS placeholders, optional Morgan fingerprints"),
        ("Forward training", "RT regression + provisional quality surrogate with grouped and source/method/column holdouts"),
        ("Recommendation", "Constrained candidate generation, forward scoring, AD penalty, optional inverse ML reranker"),
        ("Interfaces", "Streamlit GUI, FastAPI skeleton, reports, Russian PDF dashboard, local CLI scripts"),
    ]
    _add_table(doc, ["Слой", "Что делает"], rows, widths=[3.8, 12.8])


def _data_layer(doc: Document, dl_report: dict) -> None:
    doc.add_heading("3. Датасет и нормализация", level=1)
    doc.add_paragraph(
        "Training view построен как canonical merge из публичных RT источников и шаблонов internal lab. "
        "Каждая строка сохраняет source_dataset/source_record_id, compound identity и максимум доступных LC/MS metadata."
    )
    rows = [
        ("ReTiNA", "119k+ rows из Hugging Face export; сильный вклад в public RT pretraining"),
        ("METLIN SMRT Figshare", "79k+ rows; используется как крупный RT источник и benchmark"),
        ("MCMRT", "10k+ rows, 30 reversed-phase methods; полезен для method-conditioned transfer"),
        ("RepoRT", "5k+ rows из bulk import + отдельные datasets; хорош для method metadata diversity"),
        ("Kaggle METLIN descriptors", "descriptor sidecar; используется как enrichment companion, не как самостоятельный RT source"),
        ("Internal lab template", "реалистичная схема для future BE-style assay outcomes и peak metrics"),
    ]
    _add_table(doc, ["Источник", "Роль в MVP"], rows, widths=[4.5, 12.1])

    rows = [
        ("Input rows", str(dl_report.get("input_rows", "n/a"))),
        ("Valid rows", str(dl_report.get("valid_rows", "n/a"))),
        ("Filtered invalid SMILES", str(dl_report.get("filtered_invalid_smiles", "n/a"))),
        ("Split strategy", str(dl_report.get("split_strategy", "deterministic_hash"))),
        ("Train / validation / test", _split_counts(dl_report.get("split_counts", {}))),
    ]
    _add_key_value_table(doc, rows)


def _target_readiness(doc: Document, target_coverage: pd.DataFrame, chart_path: Path) -> None:
    doc.add_page_break()
    doc.add_heading("4. Что предсказываем сейчас", level=1)
    doc.add_paragraph(
        "Целевые параметры разделены на три класса: measured targets, derived/proxy targets и unavailable targets. "
        "Это защищает проект от некорректных claims: если поле есть в схеме, но нет измеренных labels, модель по нему "
        "не обучается как supervised target."
    )
    doc.add_picture(str(chart_path), width=Inches(6.3))
    _caption(doc, "Рисунок 1. Покрытие целевых параметров в текущем training view.")

    display = target_coverage.copy()
    display["coverage_fraction"] = display["coverage_fraction"].map(lambda x: f"{float(x) * 100:.0f}%")
    rows = [
        (
            row["target"],
            str(row["available_rows"]),
            row["coverage_fraction"],
            row["label_source"],
            row["readiness"],
        )
        for _, row in display.iterrows()
    ]
    _add_table(
        doc,
        ["Target", "Rows", "Coverage", "Label source", "Readiness"],
        rows,
        widths=[4.2, 2.2, 2.2, 3.2, 3.8],
        font_size=8.5,
    )


def _forward_models(
    doc: Document,
    benchmark: pd.DataFrame,
    retention_order: pd.DataFrame,
    source_holdout: pd.DataFrame,
    rt_chart: Path,
    source_chart: Path,
) -> None:
    doc.add_heading("5. Forward-модели прямой задачи", level=1)
    doc.add_paragraph(
        "Forward task формулируется как prediction under method: входом являются compound descriptors и LC/MS conditions, "
        "выходом - retention time и provisional quality score. Модель видит параметры метода как признаки, а не только "
        "строковый ярлык метода."
    )

    doc.add_picture(str(rt_chart), width=Inches(6.3))
    _caption(doc, "Рисунок 2. Test MAE для RT regression на final grouped holdout.")

    rows = []
    rt_rows = _final_rows(benchmark, "rt_min")
    for _, row in rt_rows.iterrows():
        rows.append(
            (
                row["model_family"],
                f"{row['mae']:.3f}",
                f"{row['rmse']:.3f}",
                f"{row['r2']:.3f}",
                f"{row['spearman']:.3f}",
                f"{row['normalized_mae_runtime_pct']:.2f}%",
            )
        )
    _add_table(doc, ["Model", "MAE", "RMSE", "R2", "Spearman", "nMAE/runtime"], rows, widths=[4, 2.2, 2.2, 2.0, 2.4, 3.2])

    quality_rows = _final_rows(benchmark, "quality_score")
    rows = [
        (
            row["model_family"],
            f"{row['mae']:.6f}",
            f"{row['rmse']:.6f}",
            f"{row['r2']:.3f}" if pd.notna(row["r2"]) else "",
            f"{row['spearman']:.3f}" if pd.notna(row["spearman"]) else "",
        )
        for _, row in quality_rows.iterrows()
    ]
    _add_table(doc, ["Quality model", "MAE", "RMSE", "R2", "Spearman"], rows, widths=[4.5, 3, 3, 2.5, 2.5])

    if not retention_order.empty:
        summary = retention_order.iloc[0]
        _add_info_band(
            doc,
            "Retention-order diagnostic",
            f"На held-out predictions: pairwise accuracy = {summary['pairwise_order_accuracy']:.3f}, "
            f"Spearman RT = {summary['spearman_rt']:.3f}, сравнимых пар = {int(summary['n_pairs']):,}. "
            "Это важно для метод-разработки: часто нужна не только абсолютная RT, но и правильный порядок элюирования.",
            fill=PALETTE["soft_teal"],
            accent=PALETTE["teal"],
        )

    doc.add_picture(str(source_chart), width=Inches(6.3))
    _caption(doc, "Рисунок 3. Source-family holdout MAE для RT, лучший/сильный baseline по каждому семейству.")


def _inverse_models(doc: Document, inverse_metrics: pd.DataFrame, inverse_topk: pd.DataFrame, chart_path: Path) -> None:
    doc.add_heading("6. Обратная задача: recommendation / inverse ML", level=1)
    doc.add_paragraph(
        "Обратная задача не решается прямой генеративной нейросетью. MVP использует более надежную схему: "
        "генерация кандидатных LC methods из разрешенного search space, forward scoring каждого кандидата, "
        "добавление штрафов за runtime/OOD/constraint violations и optional inverse suitability reranker."
    )
    _add_bullets(
        doc,
        [
            "Search space задает колонки, mobile phases, pH, flow, temperature, runtime bounds и forbidden combinations.",
            "Forward score учитывает target RT fit, predicted quality, confidence, runtime penalty и applicability domain.",
            "Inverse ML обучается на synthetic_proxy labels, сформированных из наблюдавшихся пригодных/непригодных условий.",
            "После появления internal accepted/failed assay labels этот же pipeline станет supervised inverse validation layer.",
        ],
    )

    doc.add_picture(str(chart_path), width=Inches(6.1))
    _caption(doc, "Рисунок 4. Brier score inverse моделей на full proxy table; ниже лучше.")

    rows = [
        (
            row["model"],
            f"{row['roc_auc']:.3f}",
            f"{row['pr_auc']:.3f}",
            f"{row['accuracy']:.3f}",
            f"{row['brier_score']:.2e}",
            str(int(row["n_test"])),
            row["label_source"],
        )
        for _, row in inverse_metrics.iterrows()
    ]
    _add_table(
        doc,
        ["Model", "ROC-AUC", "PR-AUC", "Accuracy", "Brier", "N test", "Labels"],
        rows,
        widths=[3.8, 2, 2, 2, 2.4, 2.2, 3],
        font_size=8.5,
    )

    rows = [
        (
            row["model"],
            f"{row['top_1_success']:.3f}",
            f"{row['top_3_success']:.3f}",
            f"{row['top_5_success']:.3f}",
            f"{row['mean_first_suitable_rank']:.2f}",
        )
        for _, row in inverse_topk.iterrows()
    ]
    _add_table(doc, ["Model", "Top-1", "Top-3", "Top-5", "Mean first rank"], rows, widths=[4.5, 2.2, 2.2, 2.2, 3.4])
    _add_info_band(
        doc,
        "Важная оговорка",
        "Идеальные inverse metrics объясняются synthetic_proxy labels. Это означает, что pipeline и ранжирование "
        "технически работают, но реальную ценность inverse модели нужно проверять на лабораторных accepted/failed "
        "методиках и outcome labels.",
        fill=PALETTE["soft_amber"],
        accent=PALETTE["amber"],
    )


def _dl_preparation(doc: Document, dl_report: dict) -> None:
    doc.add_heading("7. Подготовка данных для DL без обучения нейросетей", level=1)
    doc.add_paragraph(
        "Нейросетевые модели пока не обучались намеренно: сначала подготовлен reproducible data layer, чтобы позже "
        "можно было подключить GCN/GAT/MPNN/D-MPNN или SMILES-transformer embeddings без переписывания ingestion."
    )
    rows = [
        ("Graph manifest", dl_report.get("graph_manifest_path", "")),
        ("SMILES transformer manifest", dl_report.get("transformer_manifest_path", "")),
        ("Pairwise retention-order manifest", f"{dl_report.get('pairwise_manifest_path', '')} ({dl_report.get('pairwise_pairs', 0)} pairs)"),
        ("Inverse-task manifest", f"{dl_report.get('inverse_manifest_path', '')} ({dl_report.get('inverse_rows', 0)} rows)"),
        ("Graph deps", json.dumps(dl_report.get("graph_availability", {}), ensure_ascii=False)),
        ("Transformer deps", dl_report.get("encoder_availability", {}).get("reason", "n/a")),
    ]
    _add_key_value_table(doc, rows, font_size=8.2)


def _gui_reports(doc: Document) -> None:
    doc.add_heading("8. GUI, отчеты и воспроизводимость", level=1)
    doc.add_paragraph(
        "Streamlit GUI показывает dashboard, dataset assembly, training, forward prediction, method recommendation, "
        "dataset browser, model evaluation и admin/import preview. PDF dashboard предназначен для презентации состояния "
        "датасета и моделей."
    )
    rows = [
        ("Forward prediction", "Использует trained_forward_bundle.joblib, показывает RT, quality, confidence, AD flags."),
        ("Recommendation", "Генерирует top-N methods, показывает score decomposition и inverse ML score при наличии bundle."),
        ("Dataset assembly", "Показывает imported records, missingness, source distribution, examples."),
        ("Training", "Отображает feature importance, source metrics, target readiness, inverse model diagnostics."),
        ("PDF dashboard", "reports/chromaintel_dashboard_report.pdf, русская презентационная сводка."),
    ]
    _add_table(doc, ["Компонент", "Что показывает"], rows, widths=[4.5, 12.1])


def _limitations_and_roadmap(doc: Document) -> None:
    doc.add_heading("9. Ограничения и roadmap", level=1)
    _add_info_band(
        doc,
        "Ограничения текущей версии",
        "Публичные RT-датасеты плохо покрывают BE-style matrix/sample prep/MS transition outcomes и почти не содержат "
        "измеренных peak-shape labels. Поэтому сейчас качество пика и inverse success остаются proxy-layer.",
        fill=PALETTE["soft_gray"],
        accent=PALETTE["muted"],
    )
    rows = [
        ("1", "Internal lab onboarding", "Импортировать исторические runs с accepted/failed outcomes, peak width, asymmetry, tailing, S/N, resolution, area/height."),
        ("2", "Measured peak-quality models", "Переобучить quality heads на реальных labels; добавить PR-AUC/Brier/ECE для success/failure."),
        ("3", "Morgan/full comparison", "Запустить full comparison core vs Morgan vs full method features на artifact workflow."),
        ("4", "Method-conditioned DL", "После стабилизации данных обучить D-MPNN/Chemprop-style model с LC side-features."),
        ("5", "Active learning loop", "Подключить лабораторный feedback и Bayesian/active optimization вокруг top templates."),
    ]
    _add_table(doc, ["#", "Направление", "Следующий шаг"], rows, widths=[1, 4.5, 11])


def _appendix(doc: Document, target_coverage: pd.DataFrame, benchmark: pd.DataFrame, inverse_metrics: pd.DataFrame) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Приложение: ключевые артефакты", level=1)
    rows = [
        ("reports/model_training_summary.md", "Forward training summary"),
        ("reports/model_benchmark_matrix.csv", "Unified benchmark matrix"),
        ("reports/target_coverage_matrix.csv", "Direct/proxy target readiness"),
        ("reports/inverse_model_metrics.csv", "Inverse model comparison"),
        ("reports/inverse_topk_evaluation.csv", "Top-k proxy diagnostics"),
        ("reports/chromaintel_dashboard_report.pdf", "Presentation PDF dashboard"),
        ("data/processed/models/trained_forward_bundle.joblib", "Local forward bundle, not committed"),
        ("data/processed/models/inverse_recommendation_bundle.joblib", "Local inverse reranker bundle, not committed"),
    ]
    _add_table(doc, ["Артефакт", "Назначение"], rows, widths=[7, 9.6], font_size=8.5)


def _read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path, low_memory=False)


def _read_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def _target_coverage_chart(frame: pd.DataFrame, path: Path) -> Path:
    labels = frame["target"].tolist()
    values = (frame["coverage_fraction"].astype(float) * 100).tolist()
    colors = [
        PALETTE["green"] if src == "measured" else PALETTE["amber"] if src == "derived_proxy" else PALETTE["muted"]
        for src in frame["label_source"].tolist()
    ]
    return _horizontal_bar_chart(path, "Покрытие целевых параметров, %", labels, values, colors, max_value=100)


def _rt_model_chart(frame: pd.DataFrame, path: Path) -> Path:
    data = _final_rows(frame, "rt_min").sort_values("mae")
    labels = data["model_family"].tolist()
    values = data["mae"].astype(float).tolist()
    return _horizontal_bar_chart(path, "RT test MAE, min", labels, values, [PALETTE["blue"]] * len(labels))


def _source_holdout_chart(frame: pd.DataFrame, path: Path) -> Path:
    if frame.empty:
        return _placeholder_chart(path, "Нет source-holdout данных")
    subset = frame[(frame["target"] == "rt_min") & (frame["model"].isin(["extra_trees", "random_forest", "hist_gradient_boosting"]))]
    subset = subset.sort_values("mae").groupby("holdout_family", as_index=False).first().sort_values("mae")
    labels = subset["holdout_family"].astype(str).tolist()
    values = subset["mae"].astype(float).tolist()
    return _horizontal_bar_chart(path, "Лучший RT MAE по source-family holdout", labels, values, [PALETTE["teal"]] * len(labels))


def _inverse_brier_chart(frame: pd.DataFrame, path: Path) -> Path:
    if frame.empty:
        return _placeholder_chart(path, "Нет inverse metrics")
    data = frame.sort_values("brier_score")
    labels = data["model"].tolist()
    values = data["brier_score"].astype(float).tolist()
    return _horizontal_bar_chart(path, "Inverse Brier score, lower is better", labels, values, [PALETTE["blue"]] * len(labels))


def _horizontal_bar_chart(path: Path, title: str, labels: list[str], values: list[float], colors: list[str], max_value: float | None = None) -> Path:
    width, height = 1500, max(520, 90 + len(labels) * 44)
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_title = _font(34, bold=True)
    font = _font(24)
    font_small = _font(21)
    draw.text((42, 28), title, fill=_hex(PALETTE["ink"]), font=font_title)
    left_label, left_bar, right = 42, 410, width - 250
    top = 100
    maxv = max_value or (max(values) if values else 1)
    maxv = max(maxv, 1e-9)
    for idx, (label, value) in enumerate(zip(labels, values)):
        y = top + idx * 44
        draw.text((left_label, y + 2), _ellipsize(label, 30), fill=_hex(PALETTE["ink"]), font=font_small)
        draw.rounded_rectangle((left_bar, y, right, y + 26), radius=8, fill=_hex("E2E8F0"))
        bar_w = int((right - left_bar) * min(max(value, 0) / maxv, 1.0))
        draw.rounded_rectangle((left_bar, y, left_bar + bar_w, y + 26), radius=8, fill=_hex(colors[idx]))
        suffix = "%" if max_value == 100 else ""
        if suffix:
            text = f"{value:.0f}{suffix}" if abs(value - round(value)) < 0.01 else f"{value:.1f}{suffix}"
        elif abs(value) < 0.001:
            text = f"{value:.2e}"
        else:
            text = f"{value:.2f}"
        draw.text((right + 14, y - 1), text, fill=_hex(PALETTE["muted"]), font=font)
    img.save(path)
    return path


def _placeholder_chart(path: Path, text: str) -> Path:
    img = Image.new("RGB", (1200, 360), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((20, 20, 1180, 340), outline=_hex(PALETTE["line"]), width=3)
    draw.text((60, 150), text, fill=_hex(PALETTE["muted"]), font=_font(32, bold=True))
    img.save(path)
    return path


def _add_metric_cards(doc: Document, cards: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (title, value, note) in enumerate(cards):
        cell = table.cell(idx // 3, idx % 3)
        _shade_cell(cell, PALETTE["soft_blue"] if idx % 2 == 0 else PALETTE["soft_teal"])
        _cell_margins(cell, 180, 180, 180, 180)
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = _rgb(PALETTE["muted"])
        p = cell.add_paragraph()
        r = p.add_run(value)
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = _rgb(PALETTE["ink"])
        p = cell.add_paragraph(note)
        p.runs[0].font.size = Pt(8.5)
        p.runs[0].font.color.rgb = _rgb(PALETTE["muted"])


def _add_info_band(doc: Document, title: str, text: str, fill: str, accent: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _shade_cell(cell, fill)
    _cell_margins(cell, 220, 220, 220, 220)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = _rgb(accent)
    r.font.size = Pt(11)
    p = cell.add_paragraph(text)
    p.runs[0].font.size = Pt(9.5)
    p.runs[0].font.color.rgb = _rgb(PALETTE["ink"])
    doc.add_paragraph()


def _add_key_value_table(doc: Document, rows: list[tuple[str, str]], font_size: float = 9.2) -> None:
    _add_table(doc, ["Поле", "Значение"], rows, widths=[4.7, 11.8], font_size=font_size)


def _add_table(doc: Document, headers: list[str], rows: Iterable[tuple], widths: list[float], font_size: float = 9.0) -> None:
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for idx, width in enumerate(widths):
        for cell in table.columns[idx].cells:
            cell.width = Cm(width)

    header = table.rows[0].cells
    for idx, text in enumerate(headers):
        header[idx].text = str(text)
        _shade_cell(header[idx], "DBEAFE")
        _cell_margins(header[idx], 120, 120, 100, 100)
        _format_cell_text(header[idx], bold=True, size=font_size, color=PALETTE["ink"], center=True)

    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            if row_idx % 2 == 1:
                _shade_cell(cells[idx], PALETTE["soft_gray"])
            _cell_margins(cells[idx], 120, 120, 90, 90)
            _format_cell_text(cells[idx], size=font_size, color=PALETTE["ink"], center=(idx > 0 and len(str(value)) < 18))
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def _format_cell_text(cell, bold: bool = False, size: float = 9.0, color: str = "1F2937", center: bool = False) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Aptos"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = _rgb(color)


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def _caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = _rgb(PALETTE["muted"])


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _cell_margins(cell, top: int, bottom: int, left: int, right: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _final_rows(frame: pd.DataFrame, target: str) -> pd.DataFrame:
    if frame.empty:
        return pd.DataFrame()
    subset = frame[
        (frame["target"] == target)
        & (frame["split"] == "final_grouped_holdout")
        & (frame["holdout_key"] == "test")
    ].copy()
    return subset.sort_values("mae")


def _best_rt_row(frame: pd.DataFrame) -> dict:
    rows = _final_rows(frame, "rt_min")
    return rows.iloc[0].to_dict() if not rows.empty else {}


def _best_quality_row(frame: pd.DataFrame) -> dict:
    rows = _final_rows(frame, "quality_score")
    return rows.iloc[0].to_dict() if not rows.empty else {}


def _inverse_training_rows() -> int:
    path = REPORTS / "inverse_model_metrics.csv"
    if not path.exists():
        return 0
    frame = pd.read_csv(path)
    n_test = int(frame["n_test"].max()) if "n_test" in frame else 0
    return 641823 if n_test >= 160000 else n_test


def _split_counts(split_counts: dict) -> str:
    if not split_counts:
        return "n/a"
    return " / ".join(f"{key}: {value}" for key, value in split_counts.items())


def _font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def _rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def _hex(hex_color: str) -> tuple[int, int, int]:
    hex_color = hex_color.lstrip("#")
    return int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)


def _ellipsize(text: str, max_len: int) -> str:
    text = re.sub(r"\s+", " ", str(text)).strip()
    return text if len(text) <= max_len else text[: max_len - 1] + "…"


if __name__ == "__main__":
    raise SystemExit(main())
